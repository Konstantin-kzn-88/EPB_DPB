from docx import Document


def _replace_in_paragraph_runs(paragraph, placeholder, value):
    """
    Аккуратно заменяет placeholder на value в paragraph, не теряя форматирование абзаца.
    Сохраняем форматирование первого run'а, в котором начинается найденный placeholder.
    """
    # Соберем склейку всех run'ов, чтобы найти индексы в едином тексте
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(run.text for run in runs)
    start = full_text.find(placeholder)
    while start != -1:
        end = start + len(placeholder)

        # Найдем соответствие "позиция в полном тексте -> (run_index, offset)"
        run_positions = []
        pos = 0
        for i, r in enumerate(runs):
            text = r.text
            run_positions.append((i, pos, pos + len(text)))  # (индекс run, начало, конец в full_text)
            pos += len(text)

        # Определим run, где начинается и заканчивается placeholder
        # start находится в промежутке [run_start, run_end)
        def locate(idx):
            for (ri, s, e) in run_positions:
                if s <= idx < e:
                    return ri, idx - s
            # если idx == длине текста, привяжем к концу последнего run
            return len(runs) - 1, len(runs[-1].text)

        start_ri, start_off = locate(start)
        end_ri, end_off = locate(end - 1)  # последняя буква плейсхолдера

        # Три части: before (в run start), middle (полностью покрытые run'ы), after (в run end)
        before = runs[start_ri].text[:start_off]
        after = runs[end_ri].text[end_off + 1:]  # +1 т.к. end_off указывает на последнюю букву плейсхолдера

        # В первый run пишем before + value + (если плейсхолдер полностью в одном run) after
        runs[start_ri].text = before + value + (after if start_ri == end_ri else "")

        # Все run'ы между началом и концом чистим
        for i in range(start_ri + 1, end_ri):
            runs[i].text = ""

        if end_ri != start_ri:
            # В конечном run оставляем только хвост после плейсхолдера
            runs[end_ri].text = after

        # Пересоберем full_text после правки и ищем следующий вхождение
        full_text = "".join(run.text for run in runs)
        start = full_text.find(placeholder, start + len(value))


def _iter_table_cells(table):
    for row in table.rows:
        for cell in row.cells:
            yield cell
            for t in cell.tables:
                yield from _iter_table_cells(t)  # вложенные таблицы


def _replace_in_container(container, placeholder, value):
    # Абзацы
    for p in container.paragraphs:
        _replace_in_paragraph_runs(p, placeholder, value)
    # Таблицы (включая вложенные)
    for t in getattr(container, "tables", []):
        for cell in _iter_table_cells(t):
            for p in cell.paragraphs:
                _replace_in_paragraph_runs(p, placeholder, value)


def fill_template(template_path, output_path, context):
    """
    Заменяет плейсхолдеры {{ VAR }} в теле документа, таблицах и колонтитулах.
    """
    doc = Document(template_path)

    def do_all_containers(placeholder, value):
        # Тело документа
        _replace_in_container(doc, placeholder, value)
        # Колонтитулы всех секций
        for section in doc.sections:
            _replace_in_container(section.header, placeholder, value)
            _replace_in_container(section.footer, placeholder, value)

    for key, val in context.items():
        placeholder = f"{{{{ {key} }}}}"  # {{ ORG_NAME }}
        do_all_containers(placeholder, val)

    doc.save(output_path)
    print(f"✅ Документ сохранён: {output_path}")


if __name__ == "__main__":
    # ==== 🔧 ДАННЫЕ ДЛЯ ЗАПОЛНЕНИЯ ====
    # context = {
    #     # --- Основная информация об экспертизе ---
    #     "EXP_NUMBER": "224",
    #     "EXP_DATE": "25.11.2025",  # дата ЭПБ
    #     "ORDER_DATE": "17.10.2025",  # дата приказа
    #     "EXP_YEAR": "2025",
    #
    #     # --- Организация ---
    #     "ORG_NAME": "Башнефть-Добыча",
    #     "ORG_LEGAL_FORM": "Общество с ограниченной ответственностью",
    #     "ORG_LEGAL_FORM_SHORT": "ООО",
    #     "ORG_INN": "0277106840",
    #     "ORG_OGRN": "1090280032699",
    #     "ORG_ADDRESS": "450052, Республика Башкортостан, г. Уфа, ул. Карла Маркса, д. 30/1",
    #     "ORG_PHONE": "+7 (347) 261-61-61",
    #     "ORG_EMAIL": "info_bn@bashneft.ru",
    #
    #     # --- Руководитель ---
    #     "DECLARATION_APPROVER": "Генеральный директор",
    #     "DECLARATION_APPROVER_FIO": "Нонява Сергей Александрович",
    #     "DECLARATION_DATE": "05.11.2025",
    #
    #     # --- ОПО ---
    #     "OPO_NAME": "Система промысловых трубопроводов Сатаевского  месторождения",
    #     "OPO_ADDRESS": "Республика Башкортостан, Сухореченский сельский совет, Бижбулякский р-н",
    #     "OPO_CLASS": "II",
    #     "REG_NUMBER": "А41-05127-0367",
    #     "OPO_SUPERIOR_ORG": "ПАО «НК «Роснефть»",
    #     "OPO_DESCRIPTION": (
    #         "Система промысловых трубопроводов Сатаевского  месторождения"
    #     ),
    #     "OPO_COMPONENTS": "0000",
    #     "OPO_SANZONE": "0000",
    #
    #     # --- Лицензии и резервы ---
    #     "LICENSE_NUMBER": "ВХ-00-015657 от 09.10.2015",
    #     "ORDER_MAT_RESERVES": (
    #         "№ 0445 от 06.04.2020 г. "
    #         "«О создании резерва материальных ресурсов для локализации и ликвидации последствий аварий "
    #         "на опасных производственных объектах»"
    #     ),
    #     "ORDER_FIN_RESERVES": (
    #         "№ 0112 от 10 февраля 2025 г. Инструкция ООО «Башнефть-Добыча» "
    #         "«Создание и использование финансового резерва для ликвидации чрезвычайных ситуаций "
    #         "природного и техногенного характера» № ПЗ-11.04И-001112 ЮЛ-305, версия 2"
    #     ),
    #     "ORDER_NASF": "№ 0127 от 02.02.2022 г.",
    #
    #     # --- НАСФ, ПАСФ, ПЧ ---
    #     "NASF_CERTIFICATE": "№16423 от 22.04.2025 г., рег.№16/2-2-477",  # номер свидетельства НАСФ
    #     "PASF_NAME": "ФГАУ «АСФ «СВПФВЧ»",  # наименование ПАСФ
    #     "PASF_CERTIFICATE": "№13307 от 01.07.2022 г., рег.№8-177",  # номер свидетельства
    #     "PASF_CONTRACT": "№БНД/У/8/1030/21/ОПБ от 28.10.2021 г.",  # договор с ПАСФ
    #     "FIRE_DEPARTMENT_NAME": "ООО «РН-Пожарная безопасность»",  # наименование ПЧ
    #     "FIRE_DEPARTMENT_CONTRACT": "БНД/У/8/1444/23/ОПБ//5702623/0774Д от 18.12.2023 г.",  # договор с ПЧ
    #
    #     # --- Объемы документов ---
    #     "DECLARATION_VOLUME": "0000",
    #     "RPZ_VOLUME": "0000",
    #     "IFL_VOLUME": "0000",
    #
    #     # --- Документы при экспертизе ---
    #     "DOCS_SUBMITTED": (
    #         "Технологический регламент ООО «Башнефть-Добыча» на эксплуатацию ОПО «Система промысловых трубопроводов Сатаевского месторождения».,"
    #         "План мероприятий по локализации и ликвидации последствий аварий ОПО «Система промысловых трубопроводов Сатаевского месторождения»"
    #     ),
    #
    #     # --- Разработчик декларации ---
    #     "DECLARATION_DEVELOPER": "ООО «Экопромпроект»",
    #     "DECLARATION_DEVELOPER_ADDRESS": "420127, Республика Татарстан, г. Казань, ул. Короленко, д. 120, офис 1",
    #     "DECLARATION_DEVELOPER_PHONE": "+7-917-286-81-25",
    #     "DECLARATION_DEVELOPER_EMAIL": "ildarkalimullin@gmail.com",
    # }
    # context = {
    #     # --- Основная информация об экспертизе ---
    #     "EXP_NUMBER": "252",
    #     "EXP_DATE": "11.03.2026",  # дата ЭПБ
    #     "ORDER_DATE": "11.02.2026",  # дата приказа
    #     "EXP_YEAR": "2026",
    #
    #     # --- Организация ---
    #     "ORG_NAME": "Оренбургнефть",
    #     "ORG_LEGAL_FORM": "Акционерное общество",
    #     "ORG_LEGAL_FORM_SHORT": "АО",
    #     "ORG_INN": "5612002469",
    #     "ORG_OGRN": "1025601802357",
    #     "ORG_ADDRESS": "461046, Оренбургская область, г.о. город Бузулук, г. Бузулук, ул. Магистральная, зд. 2",
    #     "ORG_PHONE": "8(35342) 7-36-70 ,7-70-80/8(35342) 7-32-01",
    #     "ORG_EMAIL": "orenburgneft@rosneft.ru",
    #
    #     # --- Руководитель ---
    #     "DECLARATION_APPROVER": "Генеральный директор",
    #     "DECLARATION_APPROVER_FIO": "Хлебников Сергей Павлович",
    #     "DECLARATION_DATE": "23.09.2025",
    #
    #     # --- ОПО ---
    #     "OPO_NAME": "Система промысловых трубопроводов  Волостновского месторождения",
    #     "OPO_ADDRESS": "Российская Федерация, Оренбургская область, Новосергиевский район, с/с Рыбкинский, Волостновское месторождение, Рыбкинское месторождение",
    #     "OPO_CLASS": "II",
    #     "REG_NUMBER": "А49-01497-0491",
    #     "OPO_SUPERIOR_ORG": "ПАО «НК «Роснефть»",
    #     "OPO_DESCRIPTION": (
    #         "ОПО «Система промысловых трубопроводов  Волостновского месторождения» предназначен для: транспортирования добытой нефти."
    #     ),
    #     "OPO_COMPONENTS": "Система промысловых трубопроводов  Волостновского месторождения",
    #     "OPO_SANZONE": "0000",
    #
    #     # --- Лицензии и резервы ---
    #     "LICENSE_NUMBER": "ВХ-00-017476 от 27.08.2019",
    #     "ORDER_MAT_RESERVES": (
    #         "№01097-23 от 16.02.2023 г. «О создании объектового материального резерва для "
    #         "ликвидации чрезвычайных ситуаций на объектах АО «Оренбургнефть»"
    #     ),
    #     "ORDER_FIN_RESERVES": (
    #         "№00856-25 от 4 марта  2025 г. «О создании финансового резерва для "
    #         "ликвидации чрезвычайных ситуаций на объектах АО «Оренбургнефть»"
    #     ),
    #     "ORDER_NASF": "№ 0127 от 02.02.2022г.",
    #
    #     # --- НАСФ, ПАСФ, ПЧ ---
    #     "NASF_CERTIFICATE": "№13316 от 28.07.2022 г., рег.№16/2-1-819",  # номер свидетельства НАСФ
    #     "PASF_NAME": "АО «ЦАСЭО»",  # наименование ПАСФ
    #     "PASF_CERTIFICATE": "№ 13307, рег. номер 8-177",  # номер свидетельства
    #     "PASF_CONTRACT": "№7704124/2910Д от 11.11.24 г.",  # договор с ПАСФ
    #     "FIRE_DEPARTMENT_NAME": "ООО «РН-Пожарная безопасность»",  # наименование ПЧ
    #     "FIRE_DEPARTMENT_CONTRACT": "№7704222/2917Д/5700522/1154Д от 16.12.2022 г.",  # договор с ПЧ
    #
    #     # --- Объемы документов ---
    #     "DECLARATION_VOLUME": "0000",
    #     "RPZ_VOLUME": "0000",
    #     "IFL_VOLUME": "0000",
    #
    #     # --- Документы при экспертизе ---
    #     "DOCS_SUBMITTED": (
    #         "Технологический регламент АО «Оренбургнефть» эксплуатация Системы промысловых трубопроводов Волостновского месторождения, "
    #         "План мероприятий по локализации и ликвидации последствий аварий на ОПО АО «Оренбургнефть» № 0491, Переволоцкий участок ЦЭРТ-1, "
    #         "система промысловых трубопроводов Волостновского месторождения, регистрационный номер ОПО № А49-01497-0491, Управления эксплуатации трубопроводов № П3-05 ПМЛПА-0491 ЮЛ-412"
    #     ),
    #
    #     # --- Разработчик декларации ---
    #     "DECLARATION_DEVELOPER": "ООО «Экопромпроект»",
    #     "DECLARATION_DEVELOPER_ADDRESS": "420127, Республика Татарстан, г. Казань, ул. Короленко, д. 120, офис 1",
    #     "DECLARATION_DEVELOPER_PHONE": "+7-917-286-81-25",
    #     "DECLARATION_DEVELOPER_EMAIL": "ildarkalimullin@gmail.com",
    # }


    context = {
        # --- Основная информация об экспертизе ---
        "EXP_NUMBER": "27",
        "EXP_DATE": "11.03.2026",  # дата ЭПБ
        "ORDER_DATE": "11.02.2026",  # дата приказа
        "EXP_YEAR": "2026",

        # --- Организация ---
        "ORG_NAME": "РН-Ванкор",
        "ORG_LEGAL_FORM": "Общество с ограниченной ответственностью",
        "ORG_LEGAL_FORM_SHORT": "ООО",
        "ORG_INN": "2465142996",
        "ORG_OGRN": "1162468067541",
        "ORG_ADDRESS": "660077, Красноярский край, город Красноярск, ул. 78 Добровольческой бригады, д. 15 ",
        "ORG_PHONE": "8(391) 274-56-99/8(391) 274-56-45",
        "ORG_EMAIL": "vankor@vn.rosneft.ru",

        # --- Руководитель ---
        "DECLARATION_APPROVER": "Генеральный директор",
        "DECLARATION_APPROVER_FIO": "Чернов Владимир Николаевич",
        "DECLARATION_DATE": "11.03.2026",

        # --- ОПО ---
        "OPO_NAME": "Площадка станции насосной № 2 магистрального нефтепровода",
        "OPO_ADDRESS": "Ямало-Ненецкий автономный округ, муниципальный округ Красноселькупский район, территория НПС-2, сооружение 1",
        "OPO_CLASS": "I",
        "REG_NUMBER": "А43-01109-0014",
        "OPO_SUPERIOR_ORG": "ПАО «НК «Роснефть»",
        "OPO_DESCRIPTION": (
            "Площадка станции насосной № 2 магистрального нефтепровода предназначена подогрева и подкачки нефти на магистральном нефтепроводе"
        ),
        "OPO_COMPONENTS": "Площадка станции насосной № 2 магистрального нефтепровода",
        "OPO_SANZONE": "89-196",

        # --- Лицензии и резервы ---
        "LICENSE_NUMBER": "Л057-00109-24/00634943",
        "ORDER_MAT_RESERVES": (
            "№РНВ-1039 от 16.08.2024 г. «О порядке создания, хранения, использования и восполнения резерва материальных ресурсов для ликвидации чрезвычайных ситуаций и происшествий»"
        ),
        "ORDER_FIN_RESERVES": (
            "№РНВ-1597 от 12.12.2025 г. «О создании резерва финансовых ресурсов для ликвидации чрезвычайных ситуаций и их последствий на объектах ООО «РН-Ванкор»"
        ),
        "ORDER_NASF": "№ 0127 от 02.02.2022г.",

        # --- НАСФ, ПАСФ, ПЧ ---
        "NASF_CERTIFICATE": "№19061 от 25.04.2025 г.",  # номер свидетельства НАСФ
        "PASF_NAME": "Профессиональное аварийно-спасательное формирование «Служба по ликвидации аварийных разливов нефти, проведению газоспасательных, поисково-спасательных и противофонтанных работ» ООО «РН-Ванкор»",  # наименование ПАСФ
        "PASF_CERTIFICATE": "№19062 от 25.04.2025 г.",  # номер свидетельства
        "PASF_CONTRACT": "0000",  # договор с ПАСФ
        "FIRE_DEPARTMENT_NAME": "0000",  # наименование ПЧ
        "FIRE_DEPARTMENT_CONTRACT": "0000",  # договор с ПЧ

        # --- Объемы документов ---
        "DECLARATION_VOLUME": "0000",
        "RPZ_VOLUME": "0000",
        "IFL_VOLUME": "0000",

        # --- Документы при экспертизе ---
        "DOCS_SUBMITTED": (
            "Технологический регламент ООО «РН-Ванкор»№ П1-01.05 ТР-3277 ЮЛ-583 Эксплуатация магистрального нефтепровода «Ванкорское месторождение- НПС «Пурпе», "
            "План мероприятий по локализации и ликвидации последствий аварий на опасных производственных объектах ООО «РН-Ванкор» «Площадка станции насосной № 2 магистрального нефтепровода» (№ПЗ-05 ПМЛПА-1570 ЮЛ-583, согласован 28.10.2022 г.)"
        ),

        # --- Разработчик декларации ---
        "DECLARATION_DEVELOPER": "ООО «Экопромпроект»",
        "DECLARATION_DEVELOPER_ADDRESS": "420127, Республика Татарстан, г. Казань, ул. Короленко, д. 120, офис 1",
        "DECLARATION_DEVELOPER_PHONE": "+7-917-286-81-25",
        "DECLARATION_DEVELOPER_EMAIL": "ildarkalimullin@gmail.com",
    }



    from pathlib import Path

    # Пакетные задания: тот же набор плейсхолдеров → разные шаблоны/результаты
    jobs = [
        ("template2.docx", "pismo.docx"),  # шаблон письма → итог pismo.docx
        ("template.docx", "epb.docx"),  # шаблон ЭПБ → итог epb.docx
    ]

    for tpl, out in jobs:
        if not Path(tpl).exists():
            print(f"⚠️ Шаблон не найден: {tpl}. Пропускаю.")
            continue
        fill_template(tpl, out, context)
        print(f"✅ Заполнено: {tpl} → {out}")
